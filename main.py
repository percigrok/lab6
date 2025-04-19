# main.py

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from materials import wallpapers, tiles, laminate
from docx import Document
from openpyxl import Workbook

def calculate():
    try:
        area = float(entry_area.get())
        material = combo_material.get()

        if material == "Обои":
            rolls, cost = wallpapers.calculate_wallpapers(area, 5, 500)
        elif material == "Плитка":
            tiles_needed, cost = tiles.calculate_tiles(area, 0.25, 150)
        elif material == "Ламинат":
            planks, cost = laminate.calculate_laminate(area, 0.2, 300)
        else:
            messagebox.showerror("Ошибка", "Выберите материал!")
            return

        result_text.set(f"Необходимо материала: {rolls if material == 'Обои' else tiles_needed if material == 'Плитка' else planks} шт.\nОбщая стоимость: {cost} руб.")

    except ValueError:
        messagebox.showerror("Ошибка", "Введите корректную площадь!")

def save_report():
    filetypes = [('Документ Word', '*.docx'), ('Excel файл', '*.xlsx')]
    file = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=filetypes)

    if not file:
        return

    if file.endswith('.docx'):
        doc = Document()
        doc.add_heading('Отчет о расчетах', 0)
        doc.add_paragraph(result_text.get())
        doc.save(file)
    elif file.endswith('.xlsx'):
        wb = Workbook()
        ws = wb.active
        ws.title = "Отчет"
        lines = result_text.get().split('\n')
        for i, line in enumerate(lines, 1):
            ws.cell(row=i, column=1, value=line)
        wb.save(file)

    messagebox.showinfo("Успех", "Отчет сохранен!")

root = tk.Tk()
root.title("Калькулятор отделочных материалов")

frame = ttk.Frame(root, padding=20)
frame.grid()

ttk.Label(frame, text="Площадь (кв.м):").grid(row=0, column=0, sticky="w")
entry_area = ttk.Entry(frame)
entry_area.grid(row=0, column=1)

ttk.Label(frame, text="Материал:").grid(row=1, column=0, sticky="w")
combo_material = ttk.Combobox(frame, values=["Обои", "Плитка", "Ламинат"])
combo_material.grid(row=1, column=1)

ttk.Button(frame, text="Рассчитать", command=calculate).grid(row=2, column=0, pady=10)
ttk.Button(frame, text="Сохранить отчет", command=save_report).grid(row=2, column=1)

result_text = tk.StringVar()
ttk.Label(frame, textvariable=result_text, foreground="blue").grid(row=3, column=0, columnspan=2)

root.mainloop()
